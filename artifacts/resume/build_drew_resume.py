from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name("Drew_Sepeczi_Resume.docx")

INK = RGBColor(17, 24, 39)
MUTED = RGBColor(75, 85, 99)
ACCENT = RGBColor(37, 99, 235)
RULE = "D1D5DB"


def set_font(run, name="Arial", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=0, line=1.08):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_bottom_border(paragraph, color=RULE, size="8", space="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.48)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(9.2)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.08

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.color.rgb = INK
        style.font.bold = True
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.keep_with_next = True

    styles["Heading 1"].font.size = Pt(10.4)
    styles["Heading 2"].font.size = Pt(9.6)
    styles["Heading 3"].font.size = Pt(9.4)

    for list_style in ("List Bullet", "List Bullet 2"):
        style = styles[list_style]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(8.9)
        style.paragraph_format.left_indent = Inches(0.18)
        style.paragraph_format.first_line_indent = Inches(-0.18)
        style.paragraph_format.space_after = Pt(1.5)
        style.paragraph_format.line_spacing = 1.05


def heading(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    set_paragraph_spacing(p, before=7, after=2, line=1)
    run = p.add_run(text.upper())
    set_font(run, size=10, color=ACCENT, bold=True)
    add_bottom_border(p)
    return p


def paragraph(doc, text="", size=9.15, color=INK, bold=False, italic=False, after=1.5):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=after, line=1.08)
    r = p.add_run(text)
    set_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(p, after=1.5, line=1.05)
    for part in text:
        r = p.add_run(part[0])
        set_font(r, size=8.9, color=part[1] if len(part) > 1 else INK, bold=part[2] if len(part) > 2 else False)
    return p


def role(doc, title, org, date, location=None):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=3, after=0, line=1)
    r1 = p.add_run(title)
    set_font(r1, size=9.55, color=INK, bold=True)
    r2 = p.add_run(f" | {org}")
    set_font(r2, size=9.55, color=INK, bold=True)
    r3 = p.add_run(f" | {date}")
    set_font(r3, size=9.05, color=MUTED)
    if location:
        r4 = p.add_run(f" | {location}")
        set_font(r4, size=9.05, color=MUTED)
    return p


def project(doc, name, tags, description):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=2, after=0, line=1.04)
    r1 = p.add_run(name)
    set_font(r1, size=9.1, color=INK, bold=True)
    r2 = p.add_run(f" | {tags}: ")
    set_font(r2, size=8.95, color=MUTED)
    r3 = p.add_run(description)
    set_font(r3, size=8.95, color=INK)


def add_header(doc):
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(name, after=0, line=1)
    r = name.add_run("Drew Sepeczi")
    set_font(r, size=20, color=INK, bold=True)

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(tagline, after=2, line=1)
    r = tagline.add_run("AI Product Engineer | Founder | Full-Stack Builder")
    set_font(r, size=9.8, color=ACCENT, bold=True)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(contact, after=5, line=1.03)
    contact_text = "Chicago, IL | drewsepeczi.xyz | portfolios.chat | github.com/drewsephski | linkedin.com/in/drewsepeczi | Email: add email"
    r = contact.add_run(contact_text)
    set_font(r, size=8.4, color=MUTED)


def main():
    doc = Document()
    configure_document(doc)
    add_header(doc)

    summary = (
        "AI product engineer and founder who turns rough ideas into shipped products, tools, and infrastructure. "
        "Built 30+ web projects and 10+ AI products across full-stack architecture, AI workflows, auth, billing, "
        "deployment, launch systems, and product strategy. Strongest in zero-to-one execution that holds up in production."
    )
    paragraph(doc, summary, after=3)

    heading(doc, "Core Strengths")
    strengths = [
        "AI system design, LLM workflows, RAG, agent tooling, MCP-oriented product surfaces",
        "Full-stack product architecture with React, Next.js, TypeScript, Node.js, Python, PostgreSQL, Supabase, Redis, and Vercel",
        "Founder-grade product execution: discovery, UX, rapid prototyping, launch assets, analytics, and monetization",
    ]
    for item in strengths:
        bullet(doc, [(item, INK, False)])

    heading(doc, "Experience")
    role(doc, "Founder", "Portfolios.chat", "Feb 2026 - Present")
    bullet(doc, [("Building an AI-native portfolio platform that turns resumes, GitHub, and LinkedIn profiles into conversational, recruiter-facing sites.", INK, False)])
    bullet(doc, [("Own ingestion, AI answer generation, SEO-oriented pages, deployment, launch execution, and subscription packaging.", INK, False)])

    role(doc, "Independent Consultant", "Remote", "Sep 2024 - Jan 2026")
    bullet(doc, [("Built SlotFlow, a multi-tenant event management platform with AI-powered scheduling that cut administrative overhead in half.", INK, False)])
    bullet(doc, [("Designed tenant-aware data flows, event optimization logic, and LLM recommendation pipelines for operational scheduling workflows.", INK, False)])

    role(doc, "Founder & Lead Software Engineer", "Phoenix Agency", "Jul 2023 - Aug 2024")
    bullet(doc, [("Founded a SaaS boilerplate and client MVP studio; launched NodeBase and Astra, a SaaS template that reached Top 15 on Product Hunt.", INK, False)])
    bullet(doc, [("Delivered client MVPs with reusable product architecture across auth, billing, frontend systems, database design, and deployment.", INK, False)])

    heading(doc, "Selected Products")
    project(doc, "NodeBase", "Open Source AI Visual Studio", "Visual workflow studio for composing, running, and sharing AI pipelines without locking into a closed platform.")
    project(doc, "Portfolios.chat", "AI Portfolio SaaS", "Conversational portfolio product that answers recruiter questions from resume, GitHub, and LinkedIn context.")
    project(doc, "RagBase", "RAG / Privacy", "Private document Q&A for PDFs and URLs with cited answers in-browser and no signup requirement.")
    project(doc, "Trace", "Open Source Mac App / AI Video", "AI screen recorder for founders creating launch videos, GIFs, and product demos with less editing friction.")
    project(doc, "Squido", "Open Source CLI / AI Agents", "Terminal coding agent for developers working across real codebases with read, shell, edit, and write workflows.")
    project(doc, "SquidCrawl", "Infrastructure / Scraping", "Edge-native scraper API for AI agents with one-command deployment and LLM-optimized output.")
    project(doc, "Fight Intel / ReelDiff", "Analytics + Dev Tools", "UFC analytics with AI predictions; PR walkthrough generator for AI-narrated code-change videos.")

    heading(doc, "Writing & Thought Leadership")
    paragraph(doc, "Published technical writing on production agent patterns, AI-first frontend architecture, and open agent protocols including MCP, A2A, and ACP.", size=8.95, after=0)

    heading(doc, "Technical Stack")
    paragraph(
        doc,
        "TypeScript, React, Next.js, Node.js, Python, PostgreSQL, Supabase, Redis, WebSockets, OpenAI API, Vercel AI SDK, Vercel, Edge Functions, Electron, FFmpeg, GitHub API, React Flow, auth, billing, deployment, RAG, LLM pipelines, agent tooling.",
        size=8.95,
        after=1,
    )

    doc.sections[0].footer.paragraphs[0].clear()

    doc.save(OUT)


if __name__ == "__main__":
    main()
